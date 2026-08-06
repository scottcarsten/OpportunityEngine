"""Tests for the structured tailored-résumé renderer."""

import io

import docx

from backend.documents.resume_render import (
    parse_resume_content,
    render_plain_text_preview,
    render_resume_docx,
    render_resume_pdf,
)

_PROFILE = {
    "full_name": "Jane Doe",
    "title_line": "Engineer | Consultant",
    "location": "Austin, TX",
    "phone": "555-000-1111",
    "email": "jane@example.com",
    "education": [{"institution": "State U", "degree": "B.S. Computer Science"}],
    "certifications": ["CompTIA Security+", "CompTIA Network+"],
}

_RESUME = {
    "professional_summary": "Senior engineer with a decade of infrastructure experience.",
    "core_competencies": ["Windows Server", "Active Directory", "PowerShell"],
    "experience": [
        {
            "company": "Acme Corp",
            "location": "Remote",
            "title": "IT Manager",
            "dates": "2021 - Present",
            "bullets": ["Led infrastructure planning.", "Managed Active Directory."],
        },
        {
            "company": "Old Co",
            "location": "Dallas, TX",
            "title": "Consultant",
            "dates": "2004 - 2005",
            "bullets": [],
        },
    ],
}


def test_parse_resume_content_recognizes_structured_json() -> None:
    import json

    parsed = parse_resume_content(json.dumps(_RESUME))

    assert parsed is not None
    assert parsed["professional_summary"] == _RESUME["professional_summary"]


def test_parse_resume_content_returns_none_for_legacy_plain_text() -> None:
    assert parse_resume_content("Just a plain prose résumé draft.") is None
    assert parse_resume_content("") is None


def test_plain_text_preview_includes_summary_and_compressed_role() -> None:
    preview = render_plain_text_preview(_RESUME)

    assert "Senior engineer with a decade" in preview
    assert "Old Co | Dallas, TX" in preview
    assert "Consultant | 2004 - 2005" in preview


def test_render_resume_docx_has_no_tables() -> None:
    content = render_resume_docx(_PROFILE, _RESUME)
    document = docx.Document(io.BytesIO(content))

    assert len(document.tables) == 0


def test_render_resume_docx_includes_static_and_tailored_content() -> None:
    content = render_resume_docx(_PROFILE, _RESUME)
    document = docx.Document(io.BytesIO(content))
    full_text = "\n".join(p.text for p in document.paragraphs)

    assert "Jane Doe" in full_text
    assert "State U" in full_text
    assert "CompTIA Security+" in full_text
    assert "Led infrastructure planning." in full_text


def test_render_resume_docx_compressed_role_has_no_bullets() -> None:
    content = render_resume_docx(_PROFILE, _RESUME)
    document = docx.Document(io.BytesIO(content))

    bullet_texts = [p.text for p in document.paragraphs if p.style.name == "List Bullet"]
    assert "Led infrastructure planning." in bullet_texts
    assert "Managed Active Directory." in bullet_texts
    # The compressed "Old Co" role contributed no bullet paragraphs.
    assert not any("Old Co" in text for text in bullet_texts)


def test_render_resume_pdf_produces_a_valid_pdf_signature() -> None:
    content = render_resume_pdf(_PROFILE, _RESUME)

    assert content[:5] == b"%PDF-"
    assert len(content) > 0
