"""Application-document generation tests: service logic and route, no live Opus 5 calls."""

import json
from pathlib import Path
from urllib.parse import urlparse

import pytest
from fastapi.testclient import TestClient
from sqlalchemy import delete, select, update
from sqlalchemy.exc import IntegrityError

from backend.app import create_app
from backend.config import Settings
from backend.database import Database
from backend.db.models import AuditEventRecord, GeneratedDocument, ScoreComponent, ScoringRun
from backend.documents.base import DocumentGenerationResult
from backend.models import OpportunityInput
from backend.services.constitution_service import Constitution, load_constitution
from backend.services.document_service import DocumentService
from backend.services.opportunity_service import OpportunityService
from backend.services.resume_service import ResumeService
from backend.timeutil import now_iso


class FakeDocumentProvider:
    provider_name = "fake"
    model_name = "fake-model"
    prompt_version = "test-v1"

    def __init__(
        self,
        unsupported_claims: list[str] | None = None,
        error: Exception | None = None,
    ) -> None:
        self.unsupported_claims = unsupported_claims or []
        self.error = error
        self.calls = 0
        self.last_scoring: dict | None = None

    def generate_tailored_resume(self, opportunity, master_resume, resume_bytes, constitution):
        self.calls += 1
        if self.error is not None:
            raise self.error
        resume_fields = {
            "professional_summary": f"Tailored summary for {opportunity['title']}, grounded in "
            f"{resume_bytes.decode('utf-8')}.",
            "core_competencies": ["Azure", "AWS"],
            "experience": [
                {
                    "company": "Acme Corp",
                    "location": "Remote",
                    "title": "Engineer",
                    "dates": "2021 - Present",
                    "bullets": ["Led infrastructure work."],
                },
                {
                    "company": "Old Co",
                    "location": "Dallas, TX",
                    "title": "Consultant",
                    "dates": "2004 - 2005",
                    "bullets": [],
                },
            ],
        }
        return DocumentGenerationResult(
            content=json.dumps(resume_fields),
            unsupported_claims=self.unsupported_claims,
            structured_payload={**resume_fields, "unsupported_claims": self.unsupported_claims},
        )

    def generate_cover_letter(self, opportunity, master_resume, resume_bytes, constitution):
        self.calls += 1
        if self.error is not None:
            raise self.error
        body_paragraphs = [
            f"I am writing to express interest in the {opportunity['title']} role.",
            f"My background, grounded in {resume_bytes.decode('utf-8')}, is a strong fit.",
        ]
        return DocumentGenerationResult(
            content=json.dumps({"body_paragraphs": body_paragraphs}),
            unsupported_claims=self.unsupported_claims,
            structured_payload={
                "body_paragraphs": body_paragraphs,
                "unsupported_claims": self.unsupported_claims,
            },
        )

    def generate_fit_report(self, opportunity, master_resume, resume_bytes, scoring, constitution):
        self.calls += 1
        self.last_scoring = scoring
        if self.error is not None:
            raise self.error
        return DocumentGenerationResult(
            content=f"Fit report for {opportunity['title']}: {scoring['fit_summary']}.",
            unsupported_claims=self.unsupported_claims,
            structured_payload={"unsupported_claims": self.unsupported_claims},
        )


def _services(tmp_path: Path) -> tuple[OpportunityService, ResumeService, Database, Constitution]:
    database = Database(database_path=tmp_path / "opportunity_engine.db")
    database.initialize()
    constitution = load_constitution(Path("config/constitution.json"))
    opportunity_service = OpportunityService(database, constitution)
    resume_service = ResumeService(database, constitution, storage_path=tmp_path / "resumes")
    return opportunity_service, resume_service, database, constitution


def _preparing_opportunity(service: OpportunityService) -> int:
    opportunity_id, _ = service.create_manual(
        OpportunityInput(
            title="Cloud Administrator",
            organization_name="Acme Corp",
            description="Manage Azure and AWS environments for our engineering team.",
            source_url="https://example.com/jobs/1",
            location_text="United States",
            remote_status="remote",
            engagement_type="contract",
            tax_type="1099",
            schedule_text="After hours",
            compensation_min=85,
            compensation_max=125,
            compensation_period="hour",
            requires_travel=False,
            requires_relocation=False,
            requires_clearance=False,
            replaces_full_time_work=False,
        )
    )
    service.record_review_decision(opportunity_id, "request_preparation")
    return opportunity_id


def test_generation_requires_preparing_status(tmp_path: Path) -> None:
    opp_service, resume_service, database, constitution = _services(tmp_path)
    opportunity_id, _ = opp_service.create_manual(
        OpportunityInput(
            title="Cloud Administrator",
            organization_name="Acme Corp",
            description="Manage Azure and AWS environments.",
            source_url="https://example.com/jobs/2",
            location_text="United States",
            remote_status="remote",
            engagement_type="contract",
            tax_type="1099",
            schedule_text="After hours",
            compensation_min=None,
            compensation_max=None,
            compensation_period=None,
            requires_travel=False,
            requires_relocation=False,
            requires_clearance=False,
            replaces_full_time_work=False,
        )
    )
    document_service = DocumentService(
        database, constitution, FakeDocumentProvider(), resume_service, tmp_path / "documents"
    )

    with pytest.raises(ValueError, match="preparing"):
        document_service.generate_tailored_resume(opportunity_id)


def test_generation_requires_a_master_resume(tmp_path: Path) -> None:
    opp_service, resume_service, database, constitution = _services(tmp_path)
    opportunity_id = _preparing_opportunity(opp_service)
    document_service = DocumentService(
        database, constitution, FakeDocumentProvider(), resume_service, tmp_path / "documents"
    )

    with pytest.raises(ValueError, match="master résumé"):
        document_service.generate_tailored_resume(opportunity_id)


def test_clean_generation_creates_ready_for_review_document(tmp_path: Path) -> None:
    opp_service, resume_service, database, constitution = _services(tmp_path)
    opportunity_id = _preparing_opportunity(opp_service)
    resume_service.import_master_resume("resume.txt", b"Master resume content.", "text/plain")
    document_service = DocumentService(
        database, constitution, FakeDocumentProvider(), resume_service, tmp_path / "documents"
    )

    result = document_service.generate_tailored_resume(opportunity_id)

    assert result["status"] == "ready_for_review"
    assert result["version"] == 1
    assert result["unsupported_claims"] == []

    with database.session() as session:
        docs = session.execute(select(GeneratedDocument)).scalars().all()
    assert len(docs) == 1
    assert docs[0].document_type == "tailored_resume"
    assert docs[0].status == "ready_for_review"
    assert Path(docs[0].storage_path).is_file()


def test_generation_with_unsupported_claims_is_flagged(tmp_path: Path) -> None:
    opp_service, resume_service, database, constitution = _services(tmp_path)
    opportunity_id = _preparing_opportunity(opp_service)
    resume_service.import_master_resume("resume.txt", b"Master resume content.", "text/plain")
    provider = FakeDocumentProvider(unsupported_claims=["Invented a AWS certification."])
    document_service = DocumentService(
        database, constitution, provider, resume_service, tmp_path / "documents"
    )

    result = document_service.generate_tailored_resume(opportunity_id)

    assert result["status"] == "validation_failed"
    assert result["unsupported_claims"] == ["Invented a AWS certification."]

    with database.session() as session:
        doc = session.execute(select(GeneratedDocument)).scalars().one()
    assert doc.status == "validation_failed"
    assert json.loads(doc.unsupported_claims_json) == ["Invented a AWS certification."]


def test_regenerating_creates_a_second_version_not_an_update(tmp_path: Path) -> None:
    opp_service, resume_service, database, constitution = _services(tmp_path)
    opportunity_id = _preparing_opportunity(opp_service)
    resume_service.import_master_resume("resume.txt", b"Master resume content.", "text/plain")
    document_service = DocumentService(
        database, constitution, FakeDocumentProvider(), resume_service, tmp_path / "documents"
    )

    document_service.generate_tailored_resume(opportunity_id)
    second = document_service.generate_tailored_resume(opportunity_id)

    assert second["version"] == 2
    with database.session() as session:
        docs = session.execute(select(GeneratedDocument)).scalars().all()
    assert len(docs) == 2
    assert sorted(doc.version for doc in docs) == [1, 2]


def _seed_successful_scoring(database: Database, opportunity_id: int) -> None:
    with database.session() as session:
        run = ScoringRun(
            opportunity_id=opportunity_id,
            status="succeeded",
            scoring_version="test-v1",
            provider="fake",
            model="fake-model",
            prompt_version="test-v1",
            input_hash="deadbeef",
            overall_score=82.5,
            confidence=0.9,
            fit_summary="Strong fit for this role.",
            concerns="Compensation not specified.",
            correlation_id="test-correlation",
            started_at=now_iso(),
            completed_at=now_iso(),
        )
        session.add(run)
        session.flush()
        session.add(
            ScoreComponent(
                scoring_run_id=run.id,
                component_code="skills_alignment",
                score=90.0,
                weight=0.35,
                explanation="Strong Azure/AWS overlap.",
            )
        )
        session.commit()


def test_cover_letter_generation_creates_ready_for_review_document(tmp_path: Path) -> None:
    opp_service, resume_service, database, constitution = _services(tmp_path)
    opportunity_id = _preparing_opportunity(opp_service)
    resume_service.import_master_resume("resume.txt", b"Master resume content.", "text/plain")
    document_service = DocumentService(
        database, constitution, FakeDocumentProvider(), resume_service, tmp_path / "documents"
    )

    result = document_service.generate_cover_letter(opportunity_id)

    assert result["status"] == "ready_for_review"
    with database.session() as session:
        doc = session.execute(select(GeneratedDocument)).scalars().one()
    assert doc.document_type == "cover_letter"


def test_fit_report_requires_successful_scoring(tmp_path: Path) -> None:
    opp_service, resume_service, database, constitution = _services(tmp_path)
    opportunity_id = _preparing_opportunity(opp_service)
    resume_service.import_master_resume("resume.txt", b"Master resume content.", "text/plain")
    document_service = DocumentService(
        database, constitution, FakeDocumentProvider(), resume_service, tmp_path / "documents"
    )

    with pytest.raises(ValueError, match="score this opportunity"):
        document_service.generate_fit_report(opportunity_id)


def test_fit_report_synthesizes_the_latest_successful_scoring_run(tmp_path: Path) -> None:
    opp_service, resume_service, database, constitution = _services(tmp_path)
    opportunity_id = _preparing_opportunity(opp_service)
    resume_service.import_master_resume("resume.txt", b"Master resume content.", "text/plain")
    _seed_successful_scoring(database, opportunity_id)
    provider = FakeDocumentProvider()
    document_service = DocumentService(
        database, constitution, provider, resume_service, tmp_path / "documents"
    )

    result = document_service.generate_fit_report(opportunity_id)

    assert result["status"] == "ready_for_review"
    with database.session() as session:
        doc = session.execute(select(GeneratedDocument)).scalars().one()
    assert doc.document_type == "fit_report"
    assert provider.last_scoring["overall_score"] == 82.5
    assert provider.last_scoring["fit_summary"] == "Strong fit for this role."
    assert provider.last_scoring["components"][0]["code"] == "skills_alignment"


def test_provider_failure_records_audit_event_and_no_document(tmp_path: Path) -> None:
    opp_service, resume_service, database, constitution = _services(tmp_path)
    opportunity_id = _preparing_opportunity(opp_service)
    resume_service.import_master_resume("resume.txt", b"Master resume content.", "text/plain")
    document_service = DocumentService(
        database,
        constitution,
        FakeDocumentProvider(error=RuntimeError("boom")),
        resume_service,
        tmp_path / "documents",
    )

    with pytest.raises(RuntimeError, match="boom"):
        document_service.generate_tailored_resume(opportunity_id)

    with database.session() as session:
        docs = session.execute(select(GeneratedDocument)).scalars().all()
        events = session.execute(
            select(AuditEventRecord).where(
                AuditEventRecord.event_type == "document_generation_failed"
            )
        ).scalars().all()
    assert docs == []
    assert len(events) == 1
    assert "boom" in events[0].summary


def test_approving_a_ready_for_review_document(tmp_path: Path) -> None:
    opp_service, resume_service, database, constitution = _services(tmp_path)
    opportunity_id = _preparing_opportunity(opp_service)
    resume_service.import_master_resume("resume.txt", b"Master resume content.", "text/plain")
    document_service = DocumentService(
        database, constitution, FakeDocumentProvider(), resume_service, tmp_path / "documents"
    )
    result = document_service.generate_tailored_resume(opportunity_id)

    document_service.record_approval_decision(
        result["document_id"], "approve", rationale="Reads well."
    )

    with database.session() as session:
        doc = session.execute(select(GeneratedDocument)).scalars().one()
        events = session.execute(
            select(AuditEventRecord).where(AuditEventRecord.event_type == "document_approved")
        ).scalars().all()
    assert doc.status == "approved"
    assert doc.reviewed_at is not None
    assert len(events) == 1
    assert json.loads(events[0].details_json)["rationale"] == "Reads well."


def test_rejecting_a_validation_failed_document(tmp_path: Path) -> None:
    opp_service, resume_service, database, constitution = _services(tmp_path)
    opportunity_id = _preparing_opportunity(opp_service)
    resume_service.import_master_resume("resume.txt", b"Master resume content.", "text/plain")
    provider = FakeDocumentProvider(unsupported_claims=["Invented a certification."])
    document_service = DocumentService(
        database, constitution, provider, resume_service, tmp_path / "documents"
    )
    result = document_service.generate_tailored_resume(opportunity_id)

    document_service.record_approval_decision(result["document_id"], "reject")

    with database.session() as session:
        doc = session.execute(select(GeneratedDocument)).scalars().one()
    assert doc.status == "rejected"


def test_deciding_an_already_decided_document_raises(tmp_path: Path) -> None:
    opp_service, resume_service, database, constitution = _services(tmp_path)
    opportunity_id = _preparing_opportunity(opp_service)
    resume_service.import_master_resume("resume.txt", b"Master resume content.", "text/plain")
    document_service = DocumentService(
        database, constitution, FakeDocumentProvider(), resume_service, tmp_path / "documents"
    )
    result = document_service.generate_tailored_resume(opportunity_id)
    document_service.record_approval_decision(result["document_id"], "approve")

    with pytest.raises(ValueError, match="already been decided"):
        document_service.record_approval_decision(result["document_id"], "reject")


def test_deciding_a_nonexistent_document_raises(tmp_path: Path) -> None:
    _, resume_service, database, constitution = _services(tmp_path)
    document_service = DocumentService(
        database, constitution, FakeDocumentProvider(), resume_service, tmp_path / "documents"
    )

    with pytest.raises(ValueError, match="not found"):
        document_service.record_approval_decision(999, "approve")


def test_approved_document_is_immutable_at_the_db_level(tmp_path: Path) -> None:
    opp_service, resume_service, database, constitution = _services(tmp_path)
    opportunity_id = _preparing_opportunity(opp_service)
    resume_service.import_master_resume("resume.txt", b"Master resume content.", "text/plain")
    document_service = DocumentService(
        database, constitution, FakeDocumentProvider(), resume_service, tmp_path / "documents"
    )
    result = document_service.generate_tailored_resume(opportunity_id)
    document_service.record_approval_decision(result["document_id"], "approve")

    with database.session() as session:
        with pytest.raises(IntegrityError, match="immutable"):
            session.execute(
                update(GeneratedDocument)
                .where(GeneratedDocument.id == result["document_id"])
                .values(status="rejected")
            )

    with database.session() as session:
        with pytest.raises(IntegrityError, match="append-only"):
            session.execute(
                delete(GeneratedDocument).where(GeneratedDocument.id == result["document_id"])
            )


@pytest.fixture
def client_and_app(tmp_path: Path):
    settings = Settings(
        database_path=tmp_path / "opportunity_engine.db",
        constitution_path=Path("config/constitution.json"),
        resume_storage_path=tmp_path / "resumes",
        document_storage_path=tmp_path / "documents",
        profile_path=Path("tests/fixtures/profile_sample.json"),
    )
    app = create_app(settings)
    with TestClient(app) as test_client:
        yield test_client, app


def _form(**overrides: str) -> dict[str, str]:
    values = {
        "title": "Cloud Administrator",
        "organization_name": "Acme Corp",
        "description": "Manage Azure and AWS environments for our engineering team.",
        "source_url": "https://example.com/jobs/1",
        "location_text": "United States",
        "remote_status": "remote",
        "engagement_type": "contract",
        "tax_type": "1099",
        "schedule_text": "After hours",
        "compensation_min": "85",
        "compensation_max": "125",
        "compensation_period": "hour",
        "requires_travel": "no",
        "requires_relocation": "no",
        "requires_clearance": "no",
        "replaces_full_time_work": "no",
    }
    values.update(overrides)
    return values


def test_generation_route_renders_draft_and_flags_claims(client_and_app) -> None:
    client, app = client_and_app
    app.state.document_provider = FakeDocumentProvider(
        unsupported_claims=["Invented a AWS certification."]
    )

    client.post(
        "/resume",
        files={"file": ("resume.txt", b"Master resume content.", "text/plain")},
        follow_redirects=False,
    )
    created = client.post("/opportunities", data=_form(), follow_redirects=False)
    detail_path = urlparse(created.headers["location"]).path

    client.post(f"{detail_path}/review", data={"decision": "request_preparation"})

    response = client.post(
        f"{detail_path}/documents/tailored-resume", follow_redirects=False
    )
    assert response.status_code == 303

    detail = client.get(detail_path)
    assert "Tailored r" in detail.text
    assert "Validation failed" in detail.text
    assert "Invented a AWS certification." in detail.text


def test_generation_route_rejects_non_preparing_opportunity(client_and_app) -> None:
    client, app = client_and_app
    app.state.document_provider = FakeDocumentProvider()

    created = client.post("/opportunities", data=_form(), follow_redirects=False)
    detail_path = urlparse(created.headers["location"]).path

    response = client.post(
        f"{detail_path}/documents/tailored-resume", follow_redirects=False
    )
    assert response.status_code == 422


def test_cover_letter_route_renders_draft(client_and_app) -> None:
    client, app = client_and_app
    app.state.document_provider = FakeDocumentProvider()

    client.post(
        "/resume",
        files={"file": ("resume.txt", b"Master resume content.", "text/plain")},
        follow_redirects=False,
    )
    created = client.post("/opportunities", data=_form(), follow_redirects=False)
    detail_path = urlparse(created.headers["location"]).path
    client.post(f"{detail_path}/review", data={"decision": "request_preparation"})

    response = client.post(f"{detail_path}/documents/cover-letter", follow_redirects=False)
    assert response.status_code == 303

    detail = client.get(detail_path)
    assert "Cover letter" in detail.text
    assert "Ready for review" in detail.text


def test_fit_report_route_requires_scoring_first(client_and_app) -> None:
    client, app = client_and_app
    app.state.document_provider = FakeDocumentProvider()

    client.post(
        "/resume",
        files={"file": ("resume.txt", b"Master resume content.", "text/plain")},
        follow_redirects=False,
    )
    created = client.post("/opportunities", data=_form(), follow_redirects=False)
    detail_path = urlparse(created.headers["location"]).path
    client.post(f"{detail_path}/review", data={"decision": "request_preparation"})

    response = client.post(f"{detail_path}/documents/fit-report", follow_redirects=False)
    assert response.status_code == 422


def test_decision_route_approves_and_shows_permanently(client_and_app) -> None:
    client, app = client_and_app
    app.state.document_provider = FakeDocumentProvider()

    client.post(
        "/resume",
        files={"file": ("resume.txt", b"Master resume content.", "text/plain")},
        follow_redirects=False,
    )
    created = client.post("/opportunities", data=_form(), follow_redirects=False)
    detail_path = urlparse(created.headers["location"]).path
    client.post(f"{detail_path}/review", data={"decision": "request_preparation"})
    client.post(f"{detail_path}/documents/tailored-resume", follow_redirects=False)

    with app.state.database.session() as session:
        document_id = session.execute(select(GeneratedDocument.id)).scalars().one()

    response = client.post(
        f"{detail_path}/documents/{document_id}/decision",
        data={"decision": "approve", "rationale": "Looks accurate."},
        follow_redirects=False,
    )
    assert response.status_code == 303

    detail = client.get(detail_path)
    assert "Approved" in detail.text
    assert "Looks accurate." in detail.text
    assert f"/documents/{document_id}/decision" not in detail.text


def test_decision_route_404s_on_mismatched_opportunity(client_and_app) -> None:
    client, app = client_and_app
    app.state.document_provider = FakeDocumentProvider()

    client.post(
        "/resume",
        files={"file": ("resume.txt", b"Master resume content.", "text/plain")},
        follow_redirects=False,
    )
    first = client.post("/opportunities", data=_form(), follow_redirects=False)
    first_path = urlparse(first.headers["location"]).path
    client.post(f"{first_path}/review", data={"decision": "request_preparation"})
    client.post(f"{first_path}/documents/tailored-resume", follow_redirects=False)

    second = client.post(
        "/opportunities", data=_form(title="Second Role", source_url="https://example.com/jobs/2"),
        follow_redirects=False,
    )
    second_path = urlparse(second.headers["location"]).path

    with app.state.database.session() as session:
        document_id = session.execute(select(GeneratedDocument.id)).scalars().one()

    response = client.post(
        f"{second_path}/documents/{document_id}/decision",
        data={"decision": "approve"},
        follow_redirects=False,
    )
    assert response.status_code == 404


def test_export_docx_route_returns_valid_file(client_and_app) -> None:
    client, app = client_and_app
    app.state.document_provider = FakeDocumentProvider()

    client.post(
        "/resume",
        files={"file": ("resume.txt", b"Master resume content.", "text/plain")},
        follow_redirects=False,
    )
    created = client.post("/opportunities", data=_form(), follow_redirects=False)
    detail_path = urlparse(created.headers["location"]).path
    client.post(f"{detail_path}/review", data={"decision": "request_preparation"})
    client.post(f"{detail_path}/documents/tailored-resume", follow_redirects=False)

    with app.state.database.session() as session:
        document_id = session.execute(select(GeneratedDocument.id)).scalars().one()

    response = client.get(f"{detail_path}/documents/{document_id}/export.docx")

    assert response.status_code == 200
    assert response.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert ".docx" in response.headers["content-disposition"]
    assert response.content[:4] == b"PK\x03\x04"


def test_export_pdf_route_returns_valid_file(client_and_app) -> None:
    client, app = client_and_app
    app.state.document_provider = FakeDocumentProvider()

    client.post(
        "/resume",
        files={"file": ("resume.txt", b"Master resume content.", "text/plain")},
        follow_redirects=False,
    )
    created = client.post("/opportunities", data=_form(), follow_redirects=False)
    detail_path = urlparse(created.headers["location"]).path
    client.post(f"{detail_path}/review", data={"decision": "request_preparation"})
    client.post(f"{detail_path}/documents/tailored-resume", follow_redirects=False)

    with app.state.database.session() as session:
        document_id = session.execute(select(GeneratedDocument.id)).scalars().one()

    response = client.get(f"{detail_path}/documents/{document_id}/export.pdf")

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/pdf"
    assert ".pdf" in response.headers["content-disposition"]
    assert response.content[:5] == b"%PDF-"


def test_export_route_404s_on_mismatched_opportunity(client_and_app) -> None:
    client, app = client_and_app
    app.state.document_provider = FakeDocumentProvider()

    client.post(
        "/resume",
        files={"file": ("resume.txt", b"Master resume content.", "text/plain")},
        follow_redirects=False,
    )
    first = client.post("/opportunities", data=_form(), follow_redirects=False)
    first_path = urlparse(first.headers["location"]).path
    client.post(f"{first_path}/review", data={"decision": "request_preparation"})
    client.post(f"{first_path}/documents/tailored-resume", follow_redirects=False)

    second = client.post(
        "/opportunities", data=_form(title="Second Role", source_url="https://example.com/jobs/3"),
        follow_redirects=False,
    )
    second_path = urlparse(second.headers["location"]).path

    with app.state.database.session() as session:
        document_id = session.execute(select(GeneratedDocument.id)).scalars().one()

    response = client.get(f"{second_path}/documents/{document_id}/export.docx")
    assert response.status_code == 404


def test_export_docx_reflects_structured_resume_and_static_profile(client_and_app) -> None:
    client, app = client_and_app
    app.state.document_provider = FakeDocumentProvider()

    client.post(
        "/resume",
        files={"file": ("resume.txt", b"Master resume content.", "text/plain")},
        follow_redirects=False,
    )
    created = client.post("/opportunities", data=_form(), follow_redirects=False)
    detail_path = urlparse(created.headers["location"]).path
    client.post(f"{detail_path}/review", data={"decision": "request_preparation"})
    client.post(f"{detail_path}/documents/tailored-resume", follow_redirects=False)

    with app.state.database.session() as session:
        document_id = session.execute(select(GeneratedDocument.id)).scalars().one()

    response = client.get(f"{detail_path}/documents/{document_id}/export.docx")
    assert response.status_code == 200

    import io

    import docx

    document = docx.Document(io.BytesIO(response.content))
    assert len(document.tables) == 0
    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "Test Candidate" in full_text  # static profile data, not AI-generated
    assert "Test Certification A" in full_text
    bullet_texts = [p.text for p in document.paragraphs if p.style.name == "List Bullet"]
    assert "Led infrastructure work." in bullet_texts
    assert not any("Old Co" in text for text in bullet_texts)  # compressed role has no bullets


def test_export_falls_back_for_legacy_plain_text_resume(client_and_app) -> None:
    client, app = client_and_app

    created = client.post("/opportunities", data=_form(), follow_redirects=False)
    detail_path = urlparse(created.headers["location"]).path

    with app.state.database.session() as session:
        opportunity_id = int(detail_path.rsplit("/", 1)[-1])
        row = GeneratedDocument(
            opportunity_id=opportunity_id,
            document_type="tailored_resume",
            version=1,
            status="ready_for_review",
            storage_path=None,
            content_hash=None,
            provider="fake",
            model="fake-model",
            prompt_version="legacy",
            unsupported_claims_json="[]",
        )
        session.add(row)
        session.flush()
        document_id = row.id
        (app.state.settings.document_storage_path).mkdir(parents=True, exist_ok=True)
        legacy_path = app.state.settings.document_storage_path / "legacy.txt"
        legacy_path.write_text("Just a plain prose résumé draft, pre-dating structured output.")
        session.execute(
            update(GeneratedDocument)
            .where(GeneratedDocument.id == document_id)
            .values(storage_path=str(legacy_path))
        )
        session.commit()

    response = client.get(f"{detail_path}/documents/{document_id}/export.docx")

    assert response.status_code == 200
    assert response.content[:4] == b"PK\x03\x04"


def test_export_docx_cover_letter_reflects_structured_content_and_static_data(
    client_and_app,
) -> None:
    client, app = client_and_app
    app.state.document_provider = FakeDocumentProvider()

    client.post(
        "/resume",
        files={"file": ("resume.txt", b"Master resume content.", "text/plain")},
        follow_redirects=False,
    )
    created = client.post("/opportunities", data=_form(), follow_redirects=False)
    detail_path = urlparse(created.headers["location"]).path
    client.post(f"{detail_path}/review", data={"decision": "request_preparation"})
    client.post(f"{detail_path}/documents/cover-letter", follow_redirects=False)

    with app.state.database.session() as session:
        document_id = session.execute(
            select(GeneratedDocument.id).where(
                GeneratedDocument.document_type == "cover_letter"
            )
        ).scalars().one()

    response = client.get(f"{detail_path}/documents/{document_id}/export.docx")
    assert response.status_code == 200

    import io

    import docx

    document = docx.Document(io.BytesIO(response.content))
    assert len(document.tables) == 0
    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "Test Candidate" in full_text  # static profile data, not AI-generated
    assert "Acme Corp Hiring Team" in full_text  # from opportunity data
    assert "I am writing to express interest" in full_text  # AI body content
    assert "Sincerely," in full_text


def test_export_falls_back_for_legacy_plain_text_cover_letter(client_and_app) -> None:
    client, app = client_and_app

    created = client.post("/opportunities", data=_form(), follow_redirects=False)
    detail_path = urlparse(created.headers["location"]).path

    with app.state.database.session() as session:
        opportunity_id = int(detail_path.rsplit("/", 1)[-1])
        row = GeneratedDocument(
            opportunity_id=opportunity_id,
            document_type="cover_letter",
            version=1,
            status="ready_for_review",
            storage_path=None,
            content_hash=None,
            provider="fake",
            model="fake-model",
            prompt_version="legacy",
            unsupported_claims_json="[]",
        )
        session.add(row)
        session.flush()
        document_id = row.id
        (app.state.settings.document_storage_path).mkdir(parents=True, exist_ok=True)
        legacy_path = app.state.settings.document_storage_path / "legacy_cover_letter.txt"
        legacy_path.write_text("Just a plain prose cover letter, pre-dating structured output.")
        session.execute(
            update(GeneratedDocument)
            .where(GeneratedDocument.id == document_id)
            .values(storage_path=str(legacy_path))
        )
        session.commit()

    response = client.get(f"{detail_path}/documents/{document_id}/export.docx")

    assert response.status_code == 200
    assert response.content[:4] == b"PK\x03\x04"
