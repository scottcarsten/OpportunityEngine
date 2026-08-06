"""Tests for the structured cover-letter renderer."""

import io

import docx

from backend.documents.cover_letter_render import (
    parse_cover_letter_content,
    render_cover_letter_docx,
    render_cover_letter_pdf,
    render_plain_text_preview,
)

_PROFILE = {
    "full_name": "Jane Doe",
    "title_line": "Engineer | Consultant",
    "location": "Austin, TX",
    "phone": "555-000-1111",
    "email": "jane@example.com",
    "education": [],
    "certifications": [],
}

_OPPORTUNITY = {"organization_name": "Acme Corp"}

_LETTER = {
    "body_paragraphs": [
        "I am writing to express interest in this role.",
        "My background in infrastructure aligns well with your needs.",
    ]
}


def test_parse_cover_letter_content_recognizes_structured_json() -> None:
    import json

    parsed = parse_cover_letter_content(json.dumps(_LETTER))

    assert parsed is not None
    assert parsed["body_paragraphs"] == _LETTER["body_paragraphs"]


def test_parse_cover_letter_content_returns_none_for_legacy_plain_text() -> None:
    assert parse_cover_letter_content("Just a plain prose cover letter draft.") is None
    assert parse_cover_letter_content("") is None


def test_plain_text_preview_includes_static_and_body_content() -> None:
    preview = render_plain_text_preview(_LETTER, _PROFILE, _OPPORTUNITY)

    assert "Jane Doe" in preview
    assert "Acme Corp Hiring Team" in preview
    assert "Dear Acme Corp Hiring Team," in preview
    assert "I am writing to express interest in this role." in preview
    assert "Sincerely," in preview


def test_stray_short_fragment_is_filtered_from_rendered_output() -> None:
    letter_with_fragment = {
        "body_paragraphs": [
            "I am writing to express interest in this role.",
            "seed",
        ]
    }

    preview = render_plain_text_preview(letter_with_fragment, _PROFILE, _OPPORTUNITY)
    assert "seed" not in preview
    assert "I am writing to express interest in this role." in preview

    docx_content = render_cover_letter_docx(_PROFILE, _OPPORTUNITY, letter_with_fragment)
    document = docx.Document(io.BytesIO(docx_content))
    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "seed" not in full_text


def test_recipient_line_falls_back_when_organization_missing() -> None:
    preview = render_plain_text_preview(_LETTER, _PROFILE, {"organization_name": ""})

    assert "Dear Hiring Team," in preview


def test_render_cover_letter_docx_has_no_tables_and_includes_static_content() -> None:
    content = render_cover_letter_docx(_PROFILE, _OPPORTUNITY, _LETTER)
    document = docx.Document(io.BytesIO(content))

    assert len(document.tables) == 0
    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "Jane Doe" in full_text
    assert "Acme Corp Hiring Team" in full_text
    assert "I am writing to express interest in this role." in full_text
    assert "Sincerely," in full_text


def test_render_cover_letter_pdf_produces_a_valid_pdf_signature() -> None:
    content = render_cover_letter_pdf(_PROFILE, _OPPORTUNITY, _LETTER)

    assert content[:5] == b"%PDF-"
    assert len(content) > 0
