"""Render parsed document blocks (`markdown_subset.parse`) to DOCX/PDF bytes.

Pure functions over already-parsed blocks — no filesystem or DB access.
Rendering is on-demand, not persisted alongside the `.txt` original:
it's cheap, deterministic, and involves no AI call, so there's nothing to
gain from storing a DOCX/PDF copy per version (`OE-ADR-025`).
"""

import io
from xml.sax.saxutils import escape

import docx
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

_PDF_HEADING_STYLES = {1: "Heading1", 2: "Heading2", 3: "Heading3"}


def render_docx(title: str, blocks: list[dict]) -> bytes:
    document = docx.Document()
    document.add_heading(title, level=0)
    for block in blocks:
        if block["type"] == "heading":
            paragraph = document.add_heading(level=min(block["level"], 3))
        else:
            paragraph = document.add_paragraph()
        for run in block["runs"]:
            paragraph.add_run(run["text"]).bold = run["bold"]

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_runs_to_html(runs: list[dict]) -> str:
    return "".join(
        f"<b>{escape(r['text'])}</b>" if r["bold"] else escape(r["text"]) for r in runs
    )


def render_pdf(title: str, blocks: list[dict]) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()

    flowables = [Paragraph(escape(title), styles["Title"]), Spacer(1, 12)]
    for block in blocks:
        html = _pdf_runs_to_html(block["runs"])
        if block["type"] == "heading":
            style_name = _PDF_HEADING_STYLES.get(block["level"], "Heading3")
            flowables.append(Paragraph(html, styles[style_name]))
        else:
            flowables.append(Paragraph(html, styles["BodyText"]))
        flowables.append(Spacer(1, 8))

    doc.build(flowables)
    return buffer.getvalue()
