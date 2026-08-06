"""Render the tailored-résumé structured content into Scott's actual template.

Unlike `backend/documents/export.py` (generic Markdown-subset rendering
for cover letters/fit reports), this renders a purpose-built résumé
layout: a static identity header/education/certifications from
`config/profile.json` (never AI-generated), plus the AI-tailored
`professional_summary`/`core_competencies`/`experience` from the
provider. Core Competencies renders as a plain list, not a table — the
one deliberate ATS-safety deviation from the original template
(`OE-ADR-026`).
"""

import io
import json
from xml.sax.saxutils import escape

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

_NAVY = "1F3864"
_GREEN = "1F6357"
_NAVY_RGB = RGBColor(0x1F, 0x38, 0x64)
_GREEN_RGB = RGBColor(0x1F, 0x63, 0x57)
_MARGIN_INCHES = 0.25


def parse_resume_content(content: str) -> dict | None:
    """Return the structured résumé dict, or `None` for legacy plain-text content."""
    try:
        parsed = json.loads(content)
    except (json.JSONDecodeError, TypeError):
        return None
    if not isinstance(parsed, dict) or "experience" not in parsed:
        return None
    return parsed


def render_plain_text_preview(resume: dict) -> str:
    lines = ["PROFESSIONAL SUMMARY", "", resume.get("professional_summary", ""), ""]
    lines += ["CORE COMPETENCIES", ""]
    lines += [f"- {c}" for c in resume.get("core_competencies", [])]
    lines += ["", "PROFESSIONAL EXPERIENCE", ""]
    for entry in resume.get("experience", []):
        lines.append(f"{entry['company']} | {entry['location']}")
        lines.append(f"{entry['title']} | {entry['dates']}")
        lines += [f"    - {b}" for b in entry.get("bullets", [])]
        lines.append("")
    return "\n".join(lines).strip()


# --- DOCX -------------------------------------------------------------


def _add_bottom_border(paragraph, color: str, size: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


_BODY_PT = 9.5
_BULLET_PT = 9
_TIGHT_SPACING = Pt(2)


def _set_body_style(document: docx.Document) -> None:
    normal = document.styles["Normal"]
    normal.font.size = Pt(_BODY_PT)
    bullet_style = document.styles["List Bullet"]
    bullet_style.font.size = Pt(_BULLET_PT)
    bullet_style.paragraph_format.space_after = Pt(1)


def _tighten(paragraph) -> None:
    paragraph.paragraph_format.space_after = _TIGHT_SPACING
    paragraph.paragraph_format.space_before = Pt(0)


def _docx_heading(document: docx.Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.color.rgb = _NAVY_RGB
    run.font.size = Pt(10.5)
    _add_bottom_border(paragraph, _NAVY, 6)


def render_resume_docx(profile: dict, resume: dict) -> bytes:
    document = docx.Document()
    for section in document.sections:
        section.top_margin = Inches(_MARGIN_INCHES)
        section.bottom_margin = Inches(_MARGIN_INCHES)
        section.left_margin = Inches(_MARGIN_INCHES)
        section.right_margin = Inches(_MARGIN_INCHES)
    _set_body_style(document)

    name = document.add_paragraph()
    name.alignment = 1  # center
    _tighten(name)
    name_run = name.add_run(profile["full_name"])
    name_run.bold = True
    name_run.font.size = Pt(17)
    name_run.font.color.rgb = _NAVY_RGB

    subtitle = document.add_paragraph()
    subtitle.alignment = 1
    _tighten(subtitle)
    subtitle_run = subtitle.add_run(profile["title_line"])
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = _GREEN_RGB

    contact = document.add_paragraph()
    contact.alignment = 1
    _tighten(contact)
    contact_run = contact.add_run(
        f"{profile['location']}  |  {profile['phone']}  |  {profile['email']}"
    )
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = _NAVY_RGB
    _add_bottom_border(contact, _NAVY, 18)

    _docx_heading(document, "Professional Summary")
    summary = document.add_paragraph(resume.get("professional_summary", ""))
    _tighten(summary)

    _docx_heading(document, "Core Competencies")
    for competency in resume.get("core_competencies", []):
        document.add_paragraph(competency, style="List Bullet")

    _docx_heading(document, "Professional Experience")
    for entry in resume.get("experience", []):
        company_line = document.add_paragraph()
        _tighten(company_line)
        company_run = company_line.add_run(f"{entry['company']}  |  {entry['location']}")
        company_run.bold = True
        company_run.font.size = Pt(_BODY_PT)

        title_line = document.add_paragraph()
        _tighten(title_line)
        title_run = title_line.add_run(entry["title"])
        title_run.bold = True
        title_run.italic = True
        title_run.font.size = Pt(_BODY_PT)
        dates_run = title_line.add_run(f"  |  {entry['dates']}")
        dates_run.italic = True
        dates_run.font.size = Pt(_BODY_PT)
        _add_bottom_border(title_line, "AAAAAA", 4)

        for bullet in entry.get("bullets", []):
            document.add_paragraph(bullet, style="List Bullet")

    _docx_heading(document, "Education")
    for edu in profile.get("education", []):
        edu_paragraph = document.add_paragraph(f"{edu['institution']} — {edu['degree']}")
        _tighten(edu_paragraph)

    _docx_heading(document, "Certifications")
    for certification in profile.get("certifications", []):
        document.add_paragraph(certification, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- PDF ----------------------------------------------------------------


def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    body_size = 9.5
    return {
        "name": ParagraphStyle(
            "ResumeName", parent=base["Title"], textColor=f"#{_NAVY}", alignment=TA_CENTER,
            fontSize=17, leading=19, spaceAfter=1,
        ),
        "subtitle": ParagraphStyle(
            "ResumeSubtitle", parent=base["Normal"], textColor=f"#{_GREEN}",
            alignment=TA_CENTER, fontSize=10, spaceAfter=1,
        ),
        "contact": ParagraphStyle(
            "ResumeContact", parent=base["Normal"], textColor=f"#{_NAVY}",
            alignment=TA_CENTER, fontSize=9, spaceAfter=4,
        ),
        "heading": ParagraphStyle(
            "ResumeHeading", parent=base["Heading2"], textColor=f"#{_NAVY}",
            fontSize=10.5, leading=12, spaceBefore=6, spaceAfter=2,
        ),
        "body": ParagraphStyle(
            "ResumeBody", parent=base["BodyText"], fontSize=body_size, leading=body_size + 2,
            spaceAfter=2,
        ),
        "bullet": ParagraphStyle(
            "ResumeBullet", parent=base["BodyText"], fontSize=body_size - 0.5,
            leading=body_size + 1, leftIndent=12, spaceAfter=1,
        ),
        "company": ParagraphStyle(
            "ResumeCompany", parent=base["BodyText"], fontSize=body_size, leading=body_size + 1,
            spaceAfter=0,
        ),
        "title": ParagraphStyle(
            "ResumeTitleLine", parent=base["BodyText"], fontSize=body_size, leading=body_size + 1,
            spaceAfter=1,
        ),
    }


def render_resume_pdf(profile: dict, resume: dict) -> bytes:
    buffer = io.BytesIO()
    margin = _MARGIN_INCHES * inch
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=margin,
        bottomMargin=margin,
        leftMargin=margin,
        rightMargin=margin,
    )
    styles = _pdf_styles()

    flowables: list = [
        Paragraph(escape(profile["full_name"]), styles["name"]),
        Paragraph(escape(profile["title_line"]), styles["subtitle"]),
        Paragraph(
            escape(f"{profile['location']}  |  {profile['phone']}  |  {profile['email']}"),
            styles["contact"],
        ),
        HRFlowable(width="100%", thickness=1.5, color=f"#{_NAVY}", spaceAfter=4),
    ]

    def heading(text: str) -> None:
        flowables.append(Paragraph(escape(text.upper()), styles["heading"]))
        flowables.append(HRFlowable(width="100%", thickness=0.5, color=f"#{_NAVY}", spaceAfter=2))

    heading("Professional Summary")
    flowables.append(Paragraph(escape(resume.get("professional_summary", "")), styles["body"]))

    heading("Core Competencies")
    for competency in resume.get("core_competencies", []):
        flowables.append(Paragraph(f"• {escape(competency)}", styles["bullet"]))

    heading("Professional Experience")
    for entry in resume.get("experience", []):
        flowables.append(
            Paragraph(
                f"<b>{escape(entry['company'])}  |  {escape(entry['location'])}</b>",
                styles["company"],
            )
        )
        flowables.append(
            Paragraph(
                f"<b><i>{escape(entry['title'])}</i></b>  |  <i>{escape(entry['dates'])}</i>",
                styles["title"],
            )
        )
        flowables.append(HRFlowable(width="100%", thickness=0.5, color="#AAAAAA", spaceAfter=1))
        for bullet in entry.get("bullets", []):
            flowables.append(Paragraph(f"• {escape(bullet)}", styles["bullet"]))
        flowables.append(Spacer(1, 3))

    heading("Education")
    for edu in profile.get("education", []):
        flowables.append(
            Paragraph(f"{escape(edu['institution'])} — {escape(edu['degree'])}", styles["body"])
        )

    heading("Certifications")
    for certification in profile.get("certifications", []):
        flowables.append(Paragraph(f"• {escape(certification)}", styles["bullet"]))

    doc.build(flowables)
    return buffer.getvalue()
