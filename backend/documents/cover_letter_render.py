"""Render the cover letter's structured body into a plain business-letter format.

Deliberately unstyled compared to `resume_render.py`'s branded template —
per Scott's direction, ATS-safety is the priority and a cover letter is
already just prose, so there's nothing to gain from adding color or
bold beyond the sender's name. Sender block, date, recipient line, and
closing are static/templated (`config/profile.json` + the opportunity's
own `organization_name`) — only `body_paragraphs` comes from the AI
(`OE-ADR-027`).
"""

import io
import json
from datetime import date
from xml.sax.saxutils import escape

import docx
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter as _letter_pagesize
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from backend.documents.resume_render import _MARGIN_INCHES

_BODY_PT = 10.5


def parse_cover_letter_content(content: str) -> dict | None:
    """Return the structured cover-letter dict, or `None` for legacy plain-text content."""
    try:
        parsed = json.loads(content)
    except (json.JSONDecodeError, TypeError):
        return None
    if not isinstance(parsed, dict) or "body_paragraphs" not in parsed:
        return None
    return parsed


def _recipient_line(opportunity: dict) -> str:
    organization = (opportunity.get("organization_name") or "").strip()
    return f"{organization} Hiring Team" if organization else "Hiring Team"


_MIN_PARAGRAPH_WORDS = 4


def _body_paragraphs(letter: dict) -> list[str]:
    """Drop stray fragments (e.g. a lone leftover token) the model occasionally emits.

    The raw AI response is preserved as-is in storage/audit history —
    this filter only affects what's actually rendered/displayed, per
    `OE-ADR-027`'s addendum: a fragment isn't a "claim" the existing
    `unsupported_claims` self-review would catch, so this is a second,
    independent safety net rather than a substitute for it.
    """
    return [p for p in letter.get("body_paragraphs", []) if len(p.split()) >= _MIN_PARAGRAPH_WORDS]


def render_plain_text_preview(letter: dict, profile: dict, opportunity: dict) -> str:
    lines = [
        profile["full_name"],
        f"{profile['location']}  |  {profile['phone']}  |  {profile['email']}",
        "",
        date.today().strftime("%B %-d, %Y"),
        "",
        _recipient_line(opportunity),
        "",
        f"Dear {_recipient_line(opportunity)},",
        "",
    ]
    for paragraph in _body_paragraphs(letter):
        lines.append(paragraph)
        lines.append("")
    lines += ["Sincerely,", profile["full_name"]]
    return "\n".join(lines).strip()


# --- DOCX -----------------------------------------------------------------


def render_cover_letter_docx(profile: dict, opportunity: dict, letter: dict) -> bytes:
    document = docx.Document()
    for section in document.sections:
        section.top_margin = Inches(_MARGIN_INCHES)
        section.bottom_margin = Inches(_MARGIN_INCHES)
        section.left_margin = Inches(_MARGIN_INCHES)
        section.right_margin = Inches(_MARGIN_INCHES)
    document.styles["Normal"].font.size = Pt(_BODY_PT)

    name = document.add_paragraph()
    name.add_run(profile["full_name"]).bold = True
    document.add_paragraph(f"{profile['location']}  |  {profile['phone']}  |  {profile['email']}")
    document.add_paragraph("")
    document.add_paragraph(date.today().strftime("%B %-d, %Y"))
    document.add_paragraph("")
    document.add_paragraph(_recipient_line(opportunity))
    document.add_paragraph("")
    document.add_paragraph(f"Dear {_recipient_line(opportunity)},")
    document.add_paragraph("")

    for paragraph in _body_paragraphs(letter):
        document.add_paragraph(paragraph)
        document.add_paragraph("")

    document.add_paragraph("Sincerely,")
    document.add_paragraph(profile["full_name"])

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- PDF --------------------------------------------------------------------


def render_cover_letter_pdf(profile: dict, opportunity: dict, letter: dict) -> bytes:
    buffer = io.BytesIO()
    margin = _MARGIN_INCHES * inch
    doc = SimpleDocTemplate(
        buffer,
        pagesize=_letter_pagesize,
        topMargin=margin,
        bottomMargin=margin,
        leftMargin=margin,
        rightMargin=margin,
    )
    base = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "CoverLetterBody", parent=base["BodyText"], fontSize=_BODY_PT,
        leading=_BODY_PT + 3, spaceAfter=10,
    )
    name_style = ParagraphStyle("CoverLetterName", parent=body_style, spaceAfter=2)

    flowables: list = [
        Paragraph(f"<b>{escape(profile['full_name'])}</b>", name_style),
        Paragraph(
            escape(f"{profile['location']}  |  {profile['phone']}  |  {profile['email']}"),
            body_style,
        ),
        Spacer(1, 10),
        Paragraph(escape(date.today().strftime("%B %-d, %Y")), body_style),
        Spacer(1, 10),
        Paragraph(escape(_recipient_line(opportunity)), body_style),
        Spacer(1, 10),
        Paragraph(escape(f"Dear {_recipient_line(opportunity)},"), body_style),
    ]
    for paragraph in _body_paragraphs(letter):
        flowables.append(Paragraph(escape(paragraph), body_style))
    flowables.append(Spacer(1, 4))
    flowables.append(Paragraph("Sincerely,", body_style))
    flowables.append(Paragraph(escape(profile["full_name"]), body_style))

    doc.build(flowables)
    return buffer.getvalue()
